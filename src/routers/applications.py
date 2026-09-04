from fastapi import APIRouter, Request, Depends, HTTPException
from src.services.ingestion import UniversalScraperService as IngestionService
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from pydantic import BaseModel
from typing import List, Optional
import json
import tempfile

from fastapi.templating import Jinja2Templates
from src.core.database import get_db
from src.core.dependencies import get_current_user
from src.models.job import Job
from src.models.application import Application
from src.models.profile import UserProfile
from src.services.resume_tailor import ResumeTailorService

try:
    from weasyprint import HTML as WeasyHTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

router = APIRouter(prefix="/applications", tags=["applications"])
templates = Jinja2Templates(directory="templates")


class ResumeUpdate(BaseModel):
    """The user's edited/approved resume data."""
    summary: str
    skills: list
    experiences: list
    projects: list
    certifications: list = []


# ─── Status Management ───


from pydantic import BaseModel
class BulkSkipRequest(BaseModel):
    job_ids: list[int]

@router.post("/bulk_skip")
async def bulk_skip(req: BulkSkipRequest, db: Session = Depends(get_db)):
    user = get_current_user(db)
    # Get existing applications for these jobs
    existing = db.query(Application).filter(
        Application.user_id == user.id,
        Application.job_id.in_(req.job_ids)
    ).all()
    
    existing_ids = {app.job_id for app in existing}
    
    # Update existing
    for app in existing:
        app.status = "SKIPPED"
        
    # Create new for those missing
    for jid in req.job_ids:
        if jid not in existing_ids:
            app = Application(job_id=jid, user_id=user.id, status="SKIPPED")
            db.add(app)
            
    db.commit()
    return {"status": "success"}

@router.post("/{job_id}/status")
async def update_status(job_id: int, status: str, db: Session = Depends(get_db)):
    user = get_current_user(db)
    app = db.query(Application).filter(
        Application.job_id == job_id, Application.user_id == user.id
    ).first()

    if not app:
        app = Application(job_id=job_id, user_id=user.id, status=status)
        db.add(app)
    else:
        app.status = status

    db.commit()
    return {"status": "success", "new_status": status}


# ─── AI Resume Builder ───

@router.get("/{job_id}/builder", response_class=HTMLResponse)
async def builder_page(request: Request, job_id: int, db: Session = Depends(get_db)):
    """The main Resume Builder page. User reviews and edits AI suggestions here."""
    user = get_current_user(db)
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
        
    # JIT Hydration: Fetch full description if missing before rendering/tailoring
    IngestionService.hydrate_job(db, job)

    # Get or create application
    app = db.query(Application).filter(
        Application.job_id == job_id, Application.user_id == user.id
    ).first()
    if not app:
        app = Application(job_id=job_id, user_id=user.id, status="APPROVED")
        db.add(app)
        db.commit()
        db.refresh(app)

    # Get profile for header info
    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()

    # Check if we already have a saved resume draft
    existing_draft = None
    if app.ai_draft_json:
        try:
            existing_draft = json.loads(app.ai_draft_json)
        except Exception:
            pass

    context = {
        "request": request,
        "job": job,
        "application": app,
        "profile": profile,
        "existing_draft": json.dumps(existing_draft) if existing_draft else "null",
    }
    return templates.TemplateResponse(
        request=request, name="applications/builder.html", context=context
    )


@router.get("/{job_id}/generate")
async def generate_resume(job_id: int, db: Session = Depends(get_db)):
    """Run AI to generate tailored resume suggestions. Returns JSON for the builder UI."""
    user = get_current_user(db)
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
        
    # JIT Hydration: Fetch full description if missing before rendering/tailoring
    IngestionService.hydrate_job(db, job)

    # Generate tailored resume
    result = ResumeTailorService.generate_tailored_resume(db, job, user.id)

    # Save as draft
    app = db.query(Application).filter(
        Application.job_id == job_id, Application.user_id == user.id
    ).first()
    if app:
        app.ai_draft_json = json.dumps(result)
        db.commit()

    return result


@router.post("/{job_id}/save-resume")
async def save_resume(job_id: int, resume: ResumeUpdate, db: Session = Depends(get_db)):
    """Save the user's edited/approved resume."""
    user = get_current_user(db)
    app = db.query(Application).filter(
        Application.job_id == job_id, Application.user_id == user.id
    ).first()
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")

    # Merge with existing draft (keep matched/missing skills from AI)
    existing = {}
    if app.ai_draft_json:
        try:
            existing = json.loads(app.ai_draft_json)
        except Exception:
            pass

    existing.update(resume.model_dump())
    app.ai_draft_json = json.dumps(existing)
    db.commit()
    return {"status": "success"}


# ─── Download ───

@router.get("/{job_id}/download/pdf")
async def download_pdf(request: Request, job_id: int, inline: bool = False, db: Session = Depends(get_db)):
    if not WEASYPRINT_AVAILABLE:
        return JSONResponse({"error": "WeasyPrint not installed"}, status_code=500)

    user = get_current_user(db)
    job = db.query(Job).filter(Job.id == job_id).first()
    app = db.query(Application).filter(
        Application.job_id == job_id, Application.user_id == user.id
    ).first()
    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()

    resume_data = json.loads(app.ai_draft_json) if app and app.ai_draft_json else {}

    html_content = templates.TemplateResponse(
        request=request,
        name="applications/resume_print.html",
        context={"request": request, "resume": resume_data, "profile": profile, "job": job}
    ).body.decode()

    tmp_pdf = tempfile.mktemp(suffix=".pdf")
    WeasyHTML(string=html_content).write_pdf(tmp_pdf)

    filename = f"Resume_{profile.name or 'Candidate'}_{job.company}.pdf".replace(" ", "_")
    
    if inline:
        from fastapi.responses import Response
        with open(tmp_pdf, "rb") as f:
            pdf_bytes = f.read()
        return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": f'inline; filename="{filename}"'})
        
    return FileResponse(tmp_pdf, filename=filename, media_type="application/pdf")

@router.get("/{job_id}/preview/html")
async def preview_html(request: Request, job_id: int, db: Session = Depends(get_db)):
    user = get_current_user(db)
    job = db.query(Job).filter(Job.id == job_id).first()
    app = db.query(Application).filter(
        Application.job_id == job_id, Application.user_id == user.id
    ).first()
    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()

    resume_data = json.loads(app.ai_draft_json) if app and app.ai_draft_json else {}

    return templates.TemplateResponse(
        request=request,
        name="applications/resume_print.html",
        context={"request": request, "resume": resume_data, "profile": profile, "job": job}
    )

@router.get("/{job_id}/download/docx")
async def download_docx(job_id: int, db: Session = Depends(get_db)):
    if not DOCX_AVAILABLE:
        return JSONResponse({"error": "python-docx not installed"}, status_code=500)

    user = get_current_user(db)
    job = db.query(Job).filter(Job.id == job_id).first()
    app = db.query(Application).filter(
        Application.job_id == job_id, Application.user_id == user.id
    ).first()
    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()

    resume_data = json.loads(app.ai_draft_json) if app and app.ai_draft_json else {}

    doc = Document()

    # Header
    name = profile.name if profile else "Candidate"
    doc.add_heading(name, 0)
    contact_parts = []
    if profile:
        if profile.location: contact_parts.append(profile.location)
        if hasattr(profile, 'user_id'):
            from src.models.user import User
            u = db.query(User).filter(User.id == profile.user_id).first()
            if u: contact_parts.append(u.email)
        if profile.phone: contact_parts.append(profile.phone)
        if profile.linkedin_url: contact_parts.append(profile.linkedin_url)
        if profile.github_url: contact_parts.append(profile.github_url)
    doc.add_paragraph(" | ".join(contact_parts))

    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(resume_data.get("summary", ""))

    # Skills
    doc.add_heading("Technical Skills", level=1)
    for cat in resume_data.get("skills", []):
        skills_str = ", ".join(cat.get("skills", []))
        doc.add_paragraph(f"{cat.get('category', '')}: {skills_str}", style="List Bullet")

    # Experience
    doc.add_heading("Professional Experience", level=1)
    for exp in resume_data.get("experiences", []):
        doc.add_heading(
            f"{exp.get('title', '')} — {exp.get('company', '')}",
            level=2
        )
        doc.add_paragraph(
            f"{exp.get('location', '')} | {exp.get('date_range', '')}"
        )
        for bullet in exp.get("bullets", []):
            text = bullet.get("tailored_text", bullet.get("original_text", ""))
            doc.add_paragraph(text, style="List Bullet")

    # Projects
    doc.add_heading("Projects", level=1)
    for proj in resume_data.get("projects", []):
        doc.add_heading(proj.get("name", ""), level=2)
        if proj.get("tech_stack"):
            doc.add_paragraph(f"Technologies: {proj['tech_stack']}")
        for bullet in proj.get("bullets", []):
            text = bullet.get("tailored_text", bullet.get("original_text", ""))
            doc.add_paragraph(text, style="List Bullet")

    # Education
    if profile and hasattr(profile, 'education') and profile.education:
        doc.add_heading("Education", level=1)
        for edu in profile.education:
            doc.add_paragraph(
                f"{edu.degree} — {edu.institution} ({edu.date_str or ''})"
            )

    tmp_docx = tempfile.mktemp(suffix=".docx")
    doc.save(tmp_docx)

    filename = f"Resume_{name}_{job.company}.docx".replace(" ", "_")
    return FileResponse(
        tmp_docx, filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@router.post("/empty_trash")
async def empty_trash(db: Session = Depends(get_db)):
    user = get_current_user(db)
    
    # Find all skipped applications
    skipped_apps = db.query(Application).filter(Application.user_id == user.id, Application.status == "SKIPPED").all()
    job_ids = [app.job_id for app in skipped_apps]
    
    # Delete applications
    db.query(Application).filter(Application.id.in_([a.id for a in skipped_apps])).delete(synchronize_session=False)
    
    # Delete the underlying jobs if they are no longer referenced (for safety, delete all job_ids we just found, assuming 1:1)
    if job_ids:
        db.query(Job).filter(Job.id.in_(job_ids)).delete(synchronize_session=False)
        
    db.commit()
    return {"status": "success", "deleted": len(job_ids)}

class PipelineUpdate(BaseModel):
    sub_status: str

class PipelineDetails(BaseModel):
    recruiter_name: Optional[str] = None
    interview_date: Optional[str] = None
    salary_expected: Optional[str] = None
    salary_offered: Optional[str] = None
    pipeline_notes: Optional[str] = None

@router.put("/{job_id}/pipeline")
async def update_pipeline_status(job_id: int, payload: PipelineUpdate, db: Session = Depends(get_db)):
    user = get_current_user(db)
    app = db.query(Application).filter(Application.job_id == job_id, Application.user_id == user.id).first()
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")
        
    app.sub_status = payload.sub_status
    
    # Track the date for the new status
    from datetime import datetime, timezone
    now_utc = datetime.now(timezone.utc)
    if app.sub_status == "APPLIED" and not app.date_applied:
        app.date_applied = now_utc
    elif app.sub_status == "SCREENING" and not app.date_screening:
        app.date_screening = now_utc
    elif app.sub_status == "INTERVIEW" and not app.date_interview:
        app.date_interview = now_utc
    elif app.sub_status == "OFFER" and not app.date_offer:
        app.date_offer = now_utc
    elif app.sub_status == "REJECTED" and not app.date_rejected:
        app.date_rejected = now_utc
        
    db.commit()
    return {"status": "success", "sub_status": app.sub_status}

@router.put("/{job_id}/details")
async def update_pipeline_details(job_id: int, payload: PipelineDetails, db: Session = Depends(get_db)):
    user = get_current_user(db)
    app = db.query(Application).filter(Application.job_id == job_id, Application.user_id == user.id).first()
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")
    
    if payload.recruiter_name is not None: app.recruiter_name = payload.recruiter_name
    if payload.salary_expected is not None: app.salary_expected = payload.salary_expected
    if payload.salary_offered is not None: app.salary_offered = payload.salary_offered
    if payload.pipeline_notes is not None: app.pipeline_notes = payload.pipeline_notes
    
    if payload.interview_date:
        import dateutil.parser
        try:
            app.interview_date = dateutil.parser.parse(payload.interview_date)
        except Exception:
            pass
            
    db.commit()
    return {"status": "success"}

@router.get("/{job_id}/generate-follow-up")
async def generate_follow_up(job_id: int, db: Session = Depends(get_db)):
    user = get_current_user(db)
    app = db.query(Application).filter(Application.job_id == job_id, Application.user_id == user.id).first()
    job = db.query(Job).filter(Job.id == job_id).first()
    
    if not app or not job:
        raise HTTPException(status_code=404, detail="Not found")
        
    from src.services.job_analysis import call_gemini
    
    prompt = f"""You are an expert career coach drafting a follow-up email for a candidate.
Company: {job.company}
Role: {job.title}
Recruiter Name: {app.recruiter_name or 'Hiring Manager'}
Current Status: {app.sub_status}
Notes from Interview/Process: {app.pipeline_notes or 'None'}

Draft a polite, professional, and concise follow-up email (no more than 3 paragraphs). 
If the status is APPLIED, ask for an update on the application.
If the status is SCREENING or INTERVIEWING, thank them for their time and reiterate interest.
If the status is OFFER, ask about next steps or politely negotiate if the user mentioned it in the notes.

Return ONLY the email body (with a subject line at the top). Do not wrap in markdown or json.
"""
    draft = call_gemini(prompt, json_mode=False)
    return {"draft": draft}


@router.get("/{job_id}/details")
async def get_pipeline_details(job_id: int, db: Session = Depends(get_db)):
    user = get_current_user(db)
    app = db.query(Application).filter(Application.job_id == job_id, Application.user_id == user.id).first()
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")
        
    return {
        "recruiter_name": app.recruiter_name or "",
        "interview_date": app.interview_date.isoformat() if app.interview_date else "",
        "salary_expected": app.salary_expected or "",
        "salary_offered": app.salary_offered or "",
        "pipeline_notes": app.pipeline_notes or "",
        "date_applied": app.date_applied.isoformat() if app.date_applied else "",
        "date_screening": app.date_screening.isoformat() if app.date_screening else "",
        "date_interview": app.date_interview.isoformat() if app.date_interview else "",
        "date_offer": app.date_offer.isoformat() if app.date_offer else "",
        "date_rejected": app.date_rejected.isoformat() if app.date_rejected else ""
    }

@router.get("/{job_id}/generate-cover-letter")
async def generate_cover_letter(job_id: int, db: Session = Depends(get_db)):
    user = get_current_user(db)
    app = db.query(Application).filter(Application.job_id == job_id, Application.user_id == user.id).first()
    job = db.query(Job).filter(Job.id == job_id).first()
    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()
    
    if not app or not job:
        raise HTTPException(status_code=404, detail="Not found")
        
    from src.services.job_analysis import call_gemini
    
    # We'll use the tailored resume as the context for the cover letter
    import json
    resume_ctx = app.ai_draft_json if app.ai_draft_json else "{}"
    if resume_ctx and len(resume_ctx) > 3000:
        resume_ctx = resume_ctx[:3000]
        
    job_ctx = f"Company: {job.company}\nTitle: {job.title}\nDescription:\n{job.description[:3000]}"
    
    prompt = f"""You are an expert career coach writing a highly compelling, modern cover letter for a candidate.
    
CANDIDATE'S TAILORED RESUME DATA (JSON):
{resume_ctx}

JOB INFO:
{job_ctx}

INSTRUCTIONS:
1. Write a professional, punchy, and confident cover letter.
2. Address it to "Hiring Manager" if no name is available.
3. Keep it to 3-4 short paragraphs.
4. Highlight 1-2 core strengths of the candidate that perfectly align with the job description.
5. End with a strong call to action.
6. Do NOT output any markdown formatting (like ```), just the raw text of the letter.
7. Use placeholders like [Your Phone Number] if needed, but fill in the Candidate's Name and Company Name automatically based on the resume data.
"""
    draft = call_gemini(prompt, json_mode=False)
    return {"draft": draft}
