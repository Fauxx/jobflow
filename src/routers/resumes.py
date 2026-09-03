from fastapi import APIRouter, Request, Depends
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
import json
import os
import tempfile
from docx import Document

from src.core.database import get_db
from src.core.dependencies import get_current_user
from src.models.resume import MasterResume

try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

router = APIRouter(prefix="/resume", tags=["resumes"])
templates = Jinja2Templates(directory="templates")

@router.get("", response_class=HTMLResponse)
async def edit_resume(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(db)
    mr = db.query(MasterResume).filter(MasterResume.user_id == user.id).first()
    
    if mr:
        resume_data = json.loads(mr.content)
    else:
        # Fallback to local resume.json
        try:
            with open("src/resume.json", "r") as f:
                resume_data = json.load(f)
        except Exception:
            resume_data = {}
            
    return templates.TemplateResponse(request=request, name="resume_edit.html", context={"request": request, "resume": resume_data})

from src.models.job import Job
from src.models.application import Application
from src.services.job_analysis import extract_clean_jd
from src.services.resume_tailor import ResumeTailorService

@router.get("/builder/{job_id}", response_class=HTMLResponse)
async def builder(request: Request, job_id: int, db: Session = Depends(get_db)):
    user = get_current_user(db)
    job = db.query(Job).filter(Job.id == job_id).first()
    
    if not job:
        return HTMLResponse("Job not found", status_code=404)
        
    # Auto-Approve Job if not already applied/approved
    app = db.query(Application).filter(Application.job_id == job.id, Application.user_id == user.id).first()
    if not app:
        app = Application(job_id=job.id, user_id=user.id, status="APPROVED")
        db.add(app)
        db.commit()
    elif app.status == "NEW" or app.status == "SKIPPED":
        app.status = "APPROVED"
        db.commit()
        
    # Extract clean JD if not done already
    if not job.description or "[AI Extraction Pending]" in job.description:
        if job.raw_scraped_text:
            job.description = extract_clean_jd(job.raw_scraped_text)
            db.commit()
            
    # Tailor resume - Cache it to save API quota!
    if app.ai_draft_json:
        tailored_data = json.loads(app.ai_draft_json)
    else:
        tailored_data = ResumeTailorService.generate_tailored_resume(db, job, user.id)
        app.ai_draft_json = json.dumps(tailored_data)
        db.commit()
    
    return templates.TemplateResponse(request=request, name="builder.html", context={
        "request": request, 
        "job": job, 
        "tailored_resume": tailored_data
    })

@router.post("/builder/{job_id}/regenerate")
async def regenerate_builder(job_id: int, db: Session = Depends(get_db)):
    user = get_current_user(db)
    job = db.query(Job).filter(Job.id == job_id).first()
    app = db.query(Application).filter(Application.job_id == job.id, Application.user_id == user.id).first()
    
    if not job or not app:
        raise HTTPException(status_code=404, detail="Job or Application not found")
        
    tailored_data = ResumeTailorService.generate_tailored_resume(db, job, user.id)
    app.ai_draft_json = json.dumps(tailored_data)
    db.commit()
    
    return {"status": "success"}

@router.post("/builder/{job_id}/save")
async def save_builder(job_id: int, request: Request, db: Session = Depends(get_db)):
    user = get_current_user(db)
    app = db.query(Application).filter(Application.job_id == job_id, Application.user_id == user.id).first()
    
    if not app:
        return {"status": "error", "message": "Application not found"}
        
    data = await request.json()
    app.ai_draft_json = json.dumps(data)
    db.commit()
    
    return {"status": "success"}

@router.post("")
async def save_resume(request: Request, db: Session = Depends(get_db)):
    data = await request.json()
    user = get_current_user(db)
    mr = db.query(MasterResume).filter(MasterResume.user_id == user.id).first()
    
    if not mr:
        mr = MasterResume(user_id=user.id, format="json", content=json.dumps(data))
        db.add(mr)
    else:
        mr.content = json.dumps(data)
        
    db.commit()
    return {"status": "success"}

@router.get("/print", response_class=HTMLResponse)
async def print_resume(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(db)
    mr = db.query(MasterResume).filter(MasterResume.user_id == user.id).first()
    resume_data = json.loads(mr.content) if mr else {}
    return templates.TemplateResponse(request=request, name="resume_print.html", context={"request": request, "resume": resume_data})

@router.get("/download/pdf")
async def download_pdf(request: Request, db: Session = Depends(get_db)):
    if not WEASYPRINT_AVAILABLE:
        return {"error": "WeasyPrint not installed"}
        
    user = get_current_user(db)
    mr = db.query(MasterResume).filter(MasterResume.user_id == user.id).first()
    resume_data = json.loads(mr.content) if mr else {}
    
    html_content = templates.TemplateResponse("resume_print.html", {"request": request, "resume": resume_data}).body.decode()
    
    tmp_pdf = tempfile.mktemp(suffix=".pdf")
    HTML(string=html_content, base_url=str(request.base_url)).write_pdf(tmp_pdf)
    
    return FileResponse(tmp_pdf, filename="Resume.pdf", media_type="application/pdf")

@router.get("/download/docx")
async def download_docx(db: Session = Depends(get_db)):
    user = get_current_user(db)
    mr = db.query(MasterResume).filter(MasterResume.user_id == user.id).first()
    resume_data = json.loads(mr.content) if mr else {}
    
    doc = Document()
    doc.add_heading(resume_data.get("name", "Candidate"), 0)
    
    contact = f"{resume_data.get('location', '')} | {resume_data.get('email', '')} | {resume_data.get('phone', '')}"
    doc.add_paragraph(contact)
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(resume_data.get("summary", ""))
    
    doc.add_heading("Skills", level=1)
    for s in resume_data.get("skills", []):
        doc.add_paragraph(f"{s.get('category')}: {s.get('items')}", style="List Bullet")
        
    doc.add_heading("Experience", level=1)
    for e in resume_data.get("experience", []):
        doc.add_heading(f"{e.get('company')} - {e.get('role')}", level=2)
        doc.add_paragraph(e.get('date', ''))
        for b in e.get('bullets', []):
            doc.add_paragraph(b, style="List Bullet")
            
    doc.add_heading("Projects", level=1)
    for p in resume_data.get("projects", []):
        doc.add_heading(p.get('title'), level=2)
        for b in p.get('bullets', []):
            doc.add_paragraph(b, style="List Bullet")
            
    doc.add_heading("Education", level=1)
    for edu in resume_data.get("education", []):
        doc.add_heading(edu.get('school'), level=2)
        doc.add_paragraph(f"{edu.get('degree')} - {edu.get('date')}")
        
    tmp_docx = tempfile.mktemp(suffix=".docx")
    doc.save(tmp_docx)
    
    return FileResponse(tmp_docx, filename="Resume.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
